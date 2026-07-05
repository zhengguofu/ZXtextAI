"""
步骤录制器 - 每步截图 + AI 文字描述 + 视频录制

在自动化执行过程中为每个步骤自动截图并使用 AI 生成文字描述。
支持导出为 Word 文档（含截图嵌入）。
"""
import asyncio
import base64
import os
import re
import logging
import time
from datetime import datetime
from typing import Optional, Dict, List, Any
from io import BytesIO

logger = logging.getLogger('django')


def _sanitize_filename(name: str, max_len: int = 80) -> str:
    """清理文件名，移除 Windows 非法字符、URL、换行等"""
    if not name:
        return 'untitled'
    # 1. 先把所有空白字符（换行、制表符、回车等）统一替换为空格，
    #    防止换行打断 URL，导致后续正则无法完整匹配
    name = re.sub(r'[\r\n\t]+', ' ', name)
    # 2. 移除完整 URL（协议+域名+路径），此时换行已统一为空格，正则不会被换行打断
    name = re.sub(r'https?://\S+', '', name)
    # 3. 替换 Windows 非法字符（包括冒号、斜杠、反斜杠、点号作为路径分隔等）
    name = re.sub(r'[\\/*?:"<>|]+', '_', name)
    # 4. 压缩连续空格和下划线
    name = re.sub(r'[_\s]+', '_', name)
    # 5. 去除首尾特殊字符
    name = name.strip('._ -')
    # 6. 截断长度
    if len(name) > max_len:
        name = name[:max_len].rstrip('._ -')
    return name or 'untitled'


class StepRecorder:
    """自动化步骤录制器：截图 + AI 描述 + 视频片段"""

    def __init__(self, output_dir: str = None, case_name: str = ""):
        """
        Args:
            output_dir: 截图保存目录
            case_name: 用例名称（用于目录命名）
        """
        self.case_name = _sanitize_filename(case_name or "Test")
        if output_dir:
            self.output_dir = output_dir
        else:
            from django.conf import settings
            self.output_dir = os.path.join(
                settings.MEDIA_ROOT, 'ui-automation', 'step-screenshots',
                self.case_name
            )
        # 最终安全检查：确保路径中无换行/回车等非法字符
        self.output_dir = self.output_dir.replace('\n', '').replace('\r', '')
        try:
            os.makedirs(self.output_dir, exist_ok=True)
        except OSError as e:
            logger.error(f"无法创建截图目录 {self.output_dir}: {e}，回退到临时目录")
            import tempfile
            fallback_dir = os.path.join(tempfile.gettempdir(), 'zx_ui_screenshots', self.case_name)
            fallback_dir = fallback_dir.replace('\n', '').replace('\r', '')
            os.makedirs(fallback_dir, exist_ok=True)
            self.output_dir = fallback_dir

        self.steps: List[Dict[str, Any]] = []
        self.step_index = 0
        self.start_time = datetime.now()
        self.video_frames: List[bytes] = []  # 视频帧（用于合成 MP4）
        self.record_video = False

    def enable_video(self, enabled: bool = True):
        """启用视频录制"""
        self.record_video = enabled

    async def capture_step(self, page, step_description: str = "",
                           step_status: str = "executed",
                           action_type: str = "",
                           element_info: str = "") -> Dict[str, Any]:
        """
        捕获一个步骤的截图和 AI 描述。

        Args:
            page: Playwright Page 对象
            step_description: 步骤描述
            step_status: 步骤状态 (executed/passed/failed/skipped)
            action_type: 操作类型 (click/input/navigate/等)
            element_info: 元素信息

        Returns:
            步骤记录字典
        """
        self.step_index += 1
        step_num = self.step_index

        timestamp = datetime.now()
        screenshot_path = ""
        screenshot_b64 = ""
        ai_description = ""
        page_url = ""
        page_title = ""

        try:
            # 获取页面信息
            page_url = getattr(page, 'url', '') or ''
            try:
                page_title = await page.title() if hasattr(page, 'title') else ''
            except Exception:
                pass

            # 截图（兼容不同版本的API）
            # 修复 BINDINGS deserialize failed: 部分 browser-use 版本的 Page 对 format 参数序列化错误
            # 策略：优先无参调用（默认PNG）→ 再尝试 type='png' → 再尝试 format='png' → 最后 dict 形式
            screenshot_bytes = None
            last_screenshot_error = None
            for call_style in ('bare', 'type', 'format_kw', 'dict'):
                try:
                    if call_style == 'bare':
                        screenshot_bytes = await page.screenshot()
                    elif call_style == 'type':
                        # Playwright 官方参数是 type='png' 不是 format='png'
                        screenshot_bytes = await page.screenshot(type='png', full_page=False)
                    elif call_style == 'format_kw':
                        screenshot_bytes = await page.screenshot(format='png', full_page=False)
                    elif call_style == 'dict':
                        screenshot_bytes = await page.screenshot({'format': 'png', 'fullPage': False})
                    break
                except Exception as se:
                    last_screenshot_error = se
                    continue
            if screenshot_bytes is None:
                raise last_screenshot_error or RuntimeError("截图失败：所有方式均失败")
            screenshot_b64 = base64.b64encode(screenshot_bytes).decode('utf-8')

            # 保存截图文件
            screenshot_filename = f"step_{step_num:04d}_{int(timestamp.timestamp())}.png"
            screenshot_path = os.path.join(self.output_dir, screenshot_filename)
            with open(screenshot_path, 'wb') as f:
                f.write(screenshot_bytes)

            # 视频帧收集
            if self.record_video:
                self.video_frames.append(screenshot_bytes)

            # AI 描述（异步生成）
            if step_description:
                ai_description = await self._generate_step_description(
                    screenshot_b64, step_description, action_type, element_info,
                    page_url, page_title
                )

        except Exception as exc:
            logger.warning(f"步骤 {step_num} 截图/分析失败: {exc}")

        step_record = {
            'step_number': step_num,
            'timestamp': timestamp.isoformat(),
            'status': step_status,
            'action_type': action_type or '',
            'description': step_description,
            'ai_analysis': ai_description,
            'screenshot_path': screenshot_path,
            'screenshot_b64': screenshot_b64,
            'page_url': page_url,
            'page_title': page_title,
            'element_info': element_info,
            'duration_ms': 0,
        }

        self.steps.append(step_record)
        return step_record

    async def _generate_step_description(self, screenshot_b64: str,
                                          step_desc: str, action_type: str,
                                          element_info: str, page_url: str,
                                          page_title: str) -> str:
        """使用 AI 生成步骤执行描述。"""
        try:
            from apps.requirement_analysis.models import AIModelService

            context = f"URL: {page_url}\n标题: {page_title}"
            if action_type:
                context += f"\n操作类型: {action_type}"
            if element_info:
                context += f"\n操作元素: {element_info}"

            messages = [
                {
                    'role': 'system',
                    'content': (
                        '你是自动化测试分析助手。请根据步骤截图和执行信息，'
                        '用中文简要描述此步骤的执行情况（1-2句话），包括：'
                        '页面状态、操作结果、是否正常。'
                        '只输出描述文本，不要添加其他内容。'
                    )
                },
                {
                    'role': 'user',
                    'content': [
                        {'type': 'text', 'text': f'步骤: {step_desc}\n{context}\n请描述步骤执行情况:'},
                        {'type': 'image_url', 'image_url': {'url': f'data:image/png;base64,{screenshot_b64}'}},
                    ]
                },
            ]

            response, config = await AIModelService.call_with_auto_model_from_roles(
                ['browser_use_vision', 'browser_use_text', 'ai_tester'],
                messages,
                max_tokens=150,
            )

            if isinstance(response, dict):
                from apps.ui_automation.captcha_handler import _extract_response_text
                return _extract_response_text(response) or ''

            return str(response).strip() if response else ''

        except Exception as exc:
            logger.debug(f"AI 步骤描述生成失败: {exc}")
            return f"[{step_desc}] 执行完成" if step_desc else "步骤执行完成"

    def get_summary(self) -> Dict[str, Any]:
        """获取执行摘要。"""
        total = len(self.steps)
        if total == 0:
            return {
                'total_steps': 0,
                'start_time': self.start_time.isoformat(),
                'end_time': datetime.now().isoformat(),
                'duration_seconds': (datetime.now() - self.start_time).total_seconds(),
            }

        passed = sum(1 for s in self.steps if s['status'] in ('executed', 'passed'))
        failed = sum(1 for s in self.steps if s['status'] == 'failed')
        skipped = sum(1 for s in self.steps if s['status'] == 'skipped')

        return {
            'total_steps': total,
            'passed': passed,
            'failed': failed,
            'skipped': skipped,
            'pass_rate': round(passed / total * 100, 1) if total > 0 else 0,
            'start_time': self.start_time.isoformat(),
            'end_time': datetime.now().isoformat(),
            'duration_seconds': (datetime.now() - self.start_time).total_seconds(),
            'case_name': self.case_name,
            'video_frames_count': len(self.video_frames) if self.record_video else 0,
        }

    def to_dict(self) -> Dict[str, Any]:
        """导出为字典（用于报告生成）。"""
        return {
            'summary': self.get_summary(),
            'steps': self.steps,
            'case_name': self.case_name,
            'output_dir': self.output_dir,
        }

    def export_word_docx(self, output_path: str = None) -> str:
        """
        导出为 Word 文档（含截图嵌入）。

        Args:
            output_path: 输出路径，默认为 output_dir 下的 report.docx

        Returns:
            生成的 docx 文件路径
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx 未安装")
            return ""

        if not output_path:
            output_path = os.path.join(self.output_dir, f"{self.case_name}_report.docx")

        doc = Document()

        # 标题
        title = doc.add_heading(f'自动化测试报告: {self.case_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 摘要
        summary = self.get_summary()
        doc.add_heading('执行摘要', level=1)
        summary_text = (
            f"执行时间: {summary['start_time']} ~ {summary['end_time']}\n"
            f"总步骤: {summary['total_steps']}\n"
            f"通过: {summary['passed']} | 失败: {summary['failed']} | 跳过: {summary['skipped']}\n"
            f"通过率: {summary['pass_rate']}%\n"
            f"总耗时: {summary['duration_seconds']:.1f}秒"
        )
        doc.add_paragraph(summary_text)

        # 步骤详情
        doc.add_heading('执行步骤详情', level=1)

        for step in self.steps:
            # 步骤标题
            status_emoji = {
                'executed': '✅', 'passed': '✅',
                'failed': '❌', 'skipped': '⏭️'
            }.get(step['status'], '▶️')

            heading_text = f"{status_emoji} 步骤 {step['step_number']}: {step['description'] or '无描述'}"
            doc.add_heading(heading_text, level=2)

            # 元信息
            meta = doc.add_paragraph()
            meta.add_run(f"状态: {step['status']} | ").bold = True
            meta.add_run(f"时间: {step['timestamp']}")
            if step['page_url']:
                meta.add_run(f"\nURL: {step['page_url']}")
            if step['action_type']:
                meta.add_run(f" | 操作: {step['action_type']}")

            # AI 分析
            if step['ai_analysis']:
                ai_para = doc.add_paragraph()
                ai_para.add_run("AI 分析: ").bold = True
                ai_para.add_run(step['ai_analysis'])

            # 嵌入截图
            if step['screenshot_path'] and os.path.exists(step['screenshot_path']):
                try:
                    doc.add_picture(step['screenshot_path'], width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as img_err:
                    doc.add_paragraph(f"[截图加载失败: {img_err}]")

            # 分隔线
            doc.add_paragraph('─' * 60)

        # 保存
        doc.save(output_path)
        logger.info(f"Word 报告已生成: {output_path}")
        return output_path

    def export_video_mp4(self, output_path: str = None, fps: int = 10) -> str:
        """
        将收集的视频帧导出为 MP4 视频文件（需要 opencv-python）。

        Args:
            output_path: 输出路径
            fps: 帧率

        Returns:
            生成的 MP4 文件路径
        """
        if not self.video_frames:
            logger.warning("没有视频帧可导出")
            return ""

        try:
            import cv2
            import numpy as np
        except ImportError:
            logger.error("opencv-python 未安装，无法生成 MP4")
            return ""

        if not output_path:
            output_path = os.path.join(self.output_dir, f"{self.case_name}_video.mp4")

        if not self.video_frames:
            return ""

        # 读取第一帧确定尺寸
        first_frame = cv2.imdecode(
            np.frombuffer(self.video_frames[0], np.uint8),
            cv2.IMREAD_COLOR
        )
        if first_frame is None:
            return ""
        height, width = first_frame.shape[:2]

        # 创建视频写入器
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')
        out = cv2.VideoWriter(output_path, fourcc, fps, (width, height))

        for frame_data in self.video_frames:
            frame = cv2.imdecode(
                np.frombuffer(frame_data, np.uint8),
                cv2.IMREAD_COLOR
            )
            if frame is not None:
                out.write(frame)

        out.release()
        logger.info(f"MP4 视频已生成: {output_path} ({len(self.video_frames)} 帧, {fps}fps)")
        return output_path

    def get_media_urls(self) -> Dict[str, str]:
        """获取媒体文件的可访问 URL。"""
        from django.conf import settings

        base_url = getattr(settings, 'MEDIA_URL', '/media/')
        rel_dir = os.path.relpath(self.output_dir, settings.MEDIA_ROOT).replace('\\', '/')

        word_path = os.path.join(self.output_dir, f"{self.case_name}_report.docx")
        video_path = os.path.join(self.output_dir, f"{self.case_name}_video.mp4")

        result = {}
        if os.path.exists(word_path):
            result['word_docx'] = f"{base_url}{rel_dir}/{self.case_name}_report.docx"
        if os.path.exists(video_path):
            result['video_mp4'] = f"{base_url}{rel_dir}/{self.case_name}_video.mp4"

        return result

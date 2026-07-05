import logging
import json
import time
import threading
import queue
from django.http import StreamingHttpResponse

import requests
from asgiref.sync import async_to_sync
from django.shortcuts import render
from rest_framework import permissions, status, viewsets
from rest_framework.decorators import action
from rest_framework.response import Response

from apps.requirement_analysis.models import AIModelService

from .models import AssistantMessage, AssistantSession, ChatMessage, DifyConfig
from .serializers import (
    AssistantMessageSerializer,
    AssistantSessionCreateSerializer,
    AssistantSessionSerializer,
    ChatMessageSerializer,
)

logger = logging.getLogger(__name__)


class AssistantSessionViewSet(viewsets.ModelViewSet):
    """Assistant session management."""

    permission_classes = [permissions.IsAuthenticated]

    def get_serializer_class(self):
        if self.action in ['create', 'update', 'partial_update']:
            return AssistantSessionCreateSerializer
        return AssistantSessionSerializer

    def get_queryset(self):
        return AssistantSession.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    @action(detail=True, methods=['post'])
    def add_message(self, request, pk=None):
        session = self.get_object()
        serializer = AssistantMessageSerializer(data=request.data)

        if serializer.is_valid():
            serializer.save(session=session)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['get'])
    def messages(self, request, pk=None):
        session = self.get_object()
        messages = session.chat_messages.all()
        serializer = ChatMessageSerializer(messages, many=True)
        return Response(serializer.data)


class ChatViewSet(viewsets.ViewSet):
    """AI chat endpoint used by the assistant page."""

    permission_classes = [permissions.IsAuthenticated]
    chat_roles = ['ai_tester', 'writer', 'reviewer', 'browser_use_text', 'browser_use_vision']

    SYSTEM_PROMPT = """你是ZX全自动测试平台智能助手，资深软件测试专家。

职责：
- 理解用户需求并执行测试相关任务
- 生成完整的测试用例（正常流程、异常场景、边界条件）
- 编写自动化测试脚本
- 分析测试结果和Bug定位

测试用例输出格式（必须使用Markdown表格）：
| 用例编号 | 测试模块 | 前置条件 | 执行步骤 | 预期结果 | 优先级 |
|---------|---------|---------|---------|---------|--------|

回答格式要求：
1. 先输出思考过程，用[思考]和[/思考]包裹
2. 然后输出正式回答
3. 测试用例必须用Markdown表格

重要：
- 用户说什么就做什么，不要拒绝
- 记住对话历史，保持上下文连贯性
- 如果用户提供URL或具体需求，基于该需求生成测试用例
- 回答要专业、完整、可执行"""

    APPRECIATION_TEXT = """

如果你觉得我的服务对你有帮助，可以扫码赞赏支持一下：

![赞赏码](/media/zanshangma.jpg)

欢迎赞赏，觉得这个项目对你有帮助的欢迎点个Star：[https://github.com/zhengguofu](https://github.com/zhengguofu)"""

    def _should_add_appreciation(self, message):
        """判断是否应该添加赞赏信息"""
        keywords = ['你是谁', '你是', '什么模型', '模型', '介绍你自己', '你叫什么', '你的名字', '你好', 'hi', 'hello']
        message_lower = message.lower()
        for keyword in keywords:
            if keyword in message_lower:
                return True
        return False

    def _append_appreciation(self, message, content):
        """如果需要，追加赞赏信息"""
        if self._should_add_appreciation(message) and '赞赏码' not in content:
            return content + self.APPRECIATION_TEXT
        return content

    @staticmethod
    def _candidate_errors(candidates):
        return [
            {
                'id': item.get('id'),
                'role': item.get('role'),
                'model_name': item.get('model_name'),
                'base_url': item.get('base_url'),
                'errors': item.get('errors') or [],
            }
            for item in candidates
            if item.get('errors')
        ]

    def _call_dify(self, dify_config, request, session, user_message, message):
        headers = {
            'Authorization': f'Bearer {dify_config.api_key}',
            'Content-Type': 'application/json',
        }
        payload = {
            'inputs': {},
            'query': message,
            'user': str(request.user.id),
            'response_mode': 'blocking',
        }
        if session.conversation_id:
            payload['conversation_id'] = session.conversation_id

        api_url = dify_config.api_url.rstrip('/')
        response = requests.post(
            f'{api_url}/chat-messages',
            headers=headers,
            json=payload,
            timeout=60,
        )

        if response.status_code != 200:
            return Response({
                'error': f'Dify API error: {response.status_code}',
                'detail': response.text,
            }, status=status.HTTP_502_BAD_GATEWAY)

        data = response.json()
        if data.get('conversation_id') and not session.conversation_id:
            session.conversation_id = data['conversation_id']
            session.save(update_fields=['conversation_id', 'updated_at'])

        assistant_message = ChatMessage.objects.create(
            session=session,
            role='assistant',
            content=data.get('answer', ''),
            conversation_id=data.get('conversation_id'),
            message_id=data.get('message_id'),
        )

        # 估算token消耗
        total_tokens = len(message) // 3 + len(data.get('answer', '')) // 3

        return Response({
            'user_message': ChatMessageSerializer(user_message).data,
            'assistant_message': ChatMessageSerializer(assistant_message).data,
            'conversation_id': data.get('conversation_id'),
            'selection_mode': 'dify',
            'usage': {
                'total_tokens': total_tokens,
                'prompt_tokens': len(message) // 3,
                'completion_tokens': len(data.get('answer', '')) // 3
            }
        })

    def _build_context_messages(self, session, user_message_text, max_context_messages=10):
        """构建上下文消息"""
        messages = [{"role": "system", "content": self.SYSTEM_PROMPT}]
        history_messages = ChatMessage.objects.filter(
            session=session
        ).order_by('created_at')[:max_context_messages]

        for msg in history_messages:
            if msg.role == 'user':
                messages.append({"role": "user", "content": msg.content})
            else:
                messages.append({"role": "assistant", "content": msg.content})

        return messages

    @action(detail=False, methods=['post'])
    def send_message(self, request):
        session_id = request.data.get('session_id')
        message = (request.data.get('message') or '').strip()
        model_id = request.data.get('model_id')

        if not session_id or not message:
            return Response(
                {'error': 'session_id and message are required'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            session = AssistantSession.objects.get(session_id=session_id, user=request.user)
        except AssistantSession.DoesNotExist:
            return Response({'error': 'session not found'}, status=status.HTTP_404_NOT_FOUND)

        session.save(update_fields=['updated_at'])

        # 先保存用户消息
        user_message = ChatMessage.objects.create(
            session=session,
            role='user',
            content=message,
            conversation_id=session.conversation_id,
        )

        try:
            # 构建上下文（包含当前消息之前的历史）
            temp_messages = list(ChatMessage.objects.filter(session=session).order_by('created_at'))
            temp_messages = [m for m in temp_messages if m.id != user_message.id]

            messages_for_ai = [{"role": "system", "content": self.SYSTEM_PROMPT}]
            for msg in temp_messages[-10:]:  # 最多10条历史
                if msg.role == 'user':
                    messages_for_ai.append({"role": "user", "content": msg.content})
                else:
                    messages_for_ai.append({"role": "assistant", "content": msg.content})
            messages_for_ai.append({"role": "user", "content": message})

            # 调用真正的AI - 使用和其他地方一样的方式
            # 使用 async_to_sync 在线层线程池中执行，避免在 ASGI 主事件循环里嵌套 event loop
            # 同时为整体调用加 90 秒上限，确保任何情况下都能在合理时间内返回
            try:
                logger.info("正在调用AI模型...")

                async def _run_ai():
                    return await AIModelService.call_with_auto_model_from_roles(
                        self.chat_roles,
                        messages_for_ai,
                        max_tokens=4096,
                        overall_timeout=80.0,
                        per_config_timeout=60.0,
                    )

                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                    future = executor.submit(lambda: async_to_sync(_run_ai)())
                    try:
                        result, config = future.result(timeout=90)
                    except concurrent.futures.TimeoutError:
                        raise Exception('AI 调用总耗时超过 90 秒，已自动取消')

                ai_content = result['choices'][0]['message']['content']
                usage = result.get('usage', {})
                model_name = config.name
                selection_mode = 'auto'
                logger.info(f"AI调用成功: {ai_content[:100]}...")
            except Exception as e:
                logger.error(f"AI调用失败: {e}")
                ai_content = self._get_mock_response(message)
                usage = {'total_tokens': len(message) // 3 + len(ai_content) // 3}
                model_name = '智能助手'
                selection_mode = 'mock'

            ai_content = self._append_appreciation(message, ai_content)

            # 保存助手消息
            assistant_message = ChatMessage.objects.create(
                session=session,
                role='assistant',
                content=ai_content,
                conversation_id=session.conversation_id,
            )

            return Response({
                'user_message': ChatMessageSerializer(user_message).data,
                'assistant_message': ChatMessageSerializer(assistant_message).data,
                'conversation_id': session.conversation_id,
                'model': model_name,
                'selection_mode': selection_mode,
                'usage': usage
            })

        except Exception as e:
            logger.error(f"错误: {e}", exc_info=True)
            fallback_content = f"你好！我是ZX全自动测试平台智能助手。你的消息我收到了。"
            assistant_message = ChatMessage.objects.create(
                session=session,
                role='assistant',
                content=fallback_content,
            )
            return Response({
                'user_message': ChatMessageSerializer(user_message).data,
                'assistant_message': ChatMessageSerializer(assistant_message).data,
                'conversation_id': session.conversation_id,
                'model': '安全回退',
                'selection_mode': 'fallback',
                'usage': {'total_tokens': len(message) // 3 + len(fallback_content) // 3}
            })

    def _get_mock_response(self, message):
        """生成模拟回复 - 根据用户需求生成实际的测试用例"""
        message_lower = message.lower()

        if '你是谁' in message_lower or '你是' in message_lower or '什么模型' in message_lower or '模型' in message_lower:
            return '我是ZX全自动测试平台智能助手，专为测试相关任务服务。'
        elif '你好' in message_lower or 'hi' in message_lower or 'hello' in message_lower:
            return '你好！我是ZX全自动测试平台智能助手。有什么可以帮助你的吗？'
        elif '联系方式' in message_lower or '联系我们' in message_lower or '客服' in message_lower or '电话' in message_lower or '邮箱' in message_lower:
            return '你可以通过以下方式联系我们：\n- 邮箱：support@zxtextai.com\n- 客服热线：400-888-8888\n- 工作时间：周一至周五 9:00-18:00'
        elif '测试用例' in message or '用例' in message:
            if '登录' in message or 'login' in message_lower:
                return self._generate_login_test_cases(message)
            elif '购物' in message or '订单' in message or '支付' in message:
                return self._generate_order_test_cases(message)
            elif 'url' in message_lower or 'http' in message_lower or '页面' in message:
                return self._generate_url_test_cases(message)
            else:
                return self._generate_generic_test_cases(message)
        elif '代码' in message or 'python' in message_lower or '脚本' in message:
            return self._generate_code_example(message)
        elif '文档' in message or '报告' in message:
            return self._generate_document_example(message)
        elif '测试' in message_lower:
            return '关于测试，我可以帮你设计测试用例、编写自动化脚本、分析Bug等等。请告诉我具体需求！'
        else:
            return f'收到你的消息：{message[:50]}...\n\n我是ZX全自动测试平台智能助手。请配置AI模型以获取完整功能！'

    def _generate_login_test_cases(self, message):
        import re
        url_match = re.search(r'https?://[\w.-]+(?:/[\w./-]*)?', message)
        url = url_match.group(0) if url_match else '登录页面'
        
        return f"""[思考]用户需要为{url}生成登录功能测试用例。我需要覆盖正常登录、异常登录、边界条件等场景。

分析测试点：
1. 正常登录流程
2. 用户名/密码为空
3. 用户名/密码错误
4. 密码强度验证
5. 验证码处理（如有）
6. 登录状态保持

生成测试用例表格...[/思考]

## {url} 登录功能测试用例

| 用例编号 | 测试模块 | 前置条件 | 执行步骤 | 预期结果 | 优先级 |
|---------|---------|---------|---------|---------|--------|
| TC-LOG-001 | 正常登录 | 用户已注册账号 | 1. 打开登录页面\n2. 输入正确用户名\n3. 输入正确密码\n4. 点击登录按钮 | 成功登录，跳转到首页 | 高 |
| TC-LOG-002 | 用户名空 | 登录页面已打开 | 1. 留空用户名\n2. 输入正确密码\n3. 点击登录按钮 | 提示"用户名不能为空" | 高 |
| TC-LOG-003 | 密码为空 | 登录页面已打开 | 1. 输入正确用户名\n2. 留空密码\n3. 点击登录按钮 | 提示"密码不能为空" | 高 |
| TC-LOG-004 | 用户名错误 | 用户已注册账号 | 1. 输入错误用户名\n2. 输入正确密码\n3. 点击登录按钮 | 提示"用户名或密码错误" | 高 |
| TC-LOG-005 | 密码错误 | 用户已注册账号 | 1. 输入正确用户名\n2. 输入错误密码\n3. 点击登录按钮 | 提示"用户名或密码错误" | 高 |
| TC-LOG-006 | 密码强度不足 | 用户未注册 | 1. 输入新用户名\n2. 输入少于6位密码\n3. 尝试注册并登录 | 提示"密码长度不能少于6位" | 中 |
| TC-LOG-007 | 连续登录失败 | 用户已注册账号 | 1. 连续5次输入错误密码\n2. 输入正确密码\n3. 点击登录按钮 | 账号被锁定或要求输入验证码 | 中 |
| TC-LOG-008 | 登录后状态保持 | 用户已登录 | 1. 登录成功\n2. 关闭浏览器\n3. 重新打开页面 | 保持登录状态或提示重新登录 | 中 |
| TC-LOG-009 | 特殊字符用户名 | 用户已注册特殊字符账号 | 1. 输入包含特殊字符的用户名\n2. 输入正确密码\n3. 点击登录按钮 | 成功登录 | 低 |
| TC-LOG-010 | XSS攻击 | 登录页面已打开 | 1. 输入<script>alert('xss')</script>作为用户名\n2. 输入密码\n3. 点击登录按钮 | 特殊字符被过滤，不执行脚本 | 高 |

如需下载测试用例，请点击消息下方的下载按钮！"""

    def _generate_url_test_cases(self, message):
        import re
        url_match = re.search(r'https?://[\w.-]+(?:/[\w./-]*)?', message)
        url = url_match.group(0) if url_match else '目标页面'
        
        return f"""[思考]用户需要为{url}生成页面测试用例。我需要分析页面功能并覆盖核心测试场景。

分析测试点：
1. 页面加载和渲染
2. 核心功能交互
3. 表单验证
4. 异常场景处理

生成测试用例表格...[/思考]

## {url} 页面测试用例

| 用例编号 | 测试模块 | 前置条件 | 执行步骤 | 预期结果 | 优先级 |
|---------|---------|---------|---------|---------|--------|
| TC-PAGE-001 | 页面正常加载 | 网络正常 | 1. 打开{url}\n2. 等待页面加载完成 | 页面正常显示，无报错 | 高 |
| TC-PAGE-002 | 页面响应时间 | 网络正常 | 1. 打开浏览器开发者工具\n2. 访问{url}\n3. 记录加载时间 | 页面加载时间<3秒 | 中 |
| TC-PAGE-003 | 页面元素完整性 | 页面已加载 | 1. 检查页面标题\n2. 检查导航栏\n3. 检查主要内容区域\n4. 检查页脚 | 所有元素正常显示，无缺失 | 高 |
| TC-PAGE-004 | 链接有效性 | 页面已加载 | 1. 点击所有可点击链接\n2. 检查跳转结果 | 所有链接正常跳转，无404 | 中 |
| TC-PAGE-005 | 表单提交（如有） | 页面已加载 | 1. 填写表单所有必填项\n2. 点击提交按钮 | 表单提交成功，显示成功提示 | 高 |
| TC-PAGE-006 | 表单验证 | 页面已加载 | 1. 留空必填项\n2. 点击提交按钮 | 显示相应字段的错误提示 | 高 |
| TC-PAGE-007 | 响应式布局 | 页面已加载 | 1. 调整浏览器窗口大小\n2. 检查不同分辨率下的显示 | 页面自适应不同屏幕尺寸 | 中 |
| TC-PAGE-008 | 图片加载 | 页面已加载 | 1. 检查所有图片\n2. 检查图片Alt属性 | 所有图片正常显示，Alt属性完整 | 低 |
| TC-PAGE-009 | 空状态处理 | 页面已加载 | 1. 清除所有数据\n2. 检查页面显示 | 显示友好的空状态提示 | 中 |
| TC-PAGE-010 | 错误页面 | 网络正常 | 1. 访问不存在的路径\n2. 检查页面显示 | 显示友好的404错误页面 | 低 |

如需下载测试用例，请点击消息下方的下载按钮！"""

    def _generate_order_test_cases(self, message):
        return f"""[思考]用户需要生成订单/购物功能测试用例。我需要覆盖商品浏览、购物车、下单、支付等完整流程。

分析测试点：
1. 商品浏览和搜索
2. 购物车操作
3. 下单流程
4. 支付功能
5. 订单状态管理

生成测试用例表格...[/思考]

## 订单/购物功能测试用例

| 用例编号 | 测试模块 | 前置条件 | 执行步骤 | 预期结果 | 优先级 |
|---------|---------|---------|---------|---------|--------|
| TC-ORDER-001 | 商品搜索 | 商品已上架 | 1. 输入商品关键词\n2. 点击搜索按钮 | 显示相关商品列表 | 高 |
| TC-ORDER-002 | 添加购物车 | 商品已浏览 | 1. 选择商品规格\n2. 点击"加入购物车" | 商品成功加入购物车 | 高 |
| TC-ORDER-003 | 修改购物车数量 | 商品已加入购物车 | 1. 进入购物车\n2. 修改商品数量\n3. 点击更新 | 数量更新成功，价格重新计算 | 高 |
| TC-ORDER-004 | 删除购物车商品 | 商品已加入购物车 | 1. 进入购物车\n2. 点击删除按钮\n3. 确认删除 | 商品从购物车移除 | 中 |
| TC-ORDER-005 | 提交订单 | 购物车有商品 | 1. 点击"去结算"\n2. 确认收货地址\n3. 选择支付方式\n4. 提交订单 | 订单创建成功，生成订单号 | 高 |
| TC-ORDER-006 | 在线支付 | 订单已创建 | 1. 选择在线支付\n2. 完成支付流程 | 支付成功，订单状态更新 | 高 |
| TC-ORDER-007 | 库存不足 | 商品库存为0 | 1. 尝试添加库存为0的商品\n2. 提交订单 | 提示库存不足，无法下单 | 高 |
| TC-ORDER-008 | 订单取消 | 订单已创建未支付 | 1. 进入订单详情\n2. 点击取消订单\n3. 确认取消 | 订单状态变为已取消 | 中 |
| TC-ORDER-009 | 订单退款 | 订单已支付 | 1. 申请退款\n2. 填写退款原因\n3. 提交申请 | 退款申请成功，等待审核 | 中 |
| TC-ORDER-010 | 重复提交订单 | 网络正常 | 1. 快速连续点击提交按钮\n2. 检查订单列表 | 只创建一个订单，防止重复提交 | 高 |

如需下载测试用例，请点击消息下方的下载按钮！"""

    def _generate_generic_test_cases(self, message):
        return f"""[思考]用户需要生成测试用例，但需求不够明确。我需要提供通用的测试用例模板，帮助用户理解测试用例的结构和内容。

生成通用测试用例模板...[/思考]

## 通用功能测试用例模板

| 用例编号 | 测试模块 | 前置条件 | 执行步骤 | 预期结果 | 优先级 |
|---------|---------|---------|---------|---------|--------|
| TC-GEN-001 | 正常流程测试 | 系统正常运行 | 1. 进入功能页面\n2. 按照业务流程操作\n3. 完成操作 | 功能正常完成，数据正确保存 | 高 |
| TC-GEN-002 | 空值输入测试 | 表单已打开 | 1. 留空所有必填字段\n2. 点击提交按钮 | 显示相应字段的错误提示 | 高 |
| TC-GEN-003 | 边界值测试 | 表单已打开 | 1. 输入字段最小值\n2. 输入字段最大值\n3. 点击提交按钮 | 边界值正常处理，无报错 | 中 |
| TC-GEN-004 | 特殊字符测试 | 表单已打开 | 1. 输入特殊字符（如<>&'"）\n2. 点击提交按钮 | 特殊字符被正确处理，无XSS风险 | 高 |
| TC-GEN-005 | 异常数据测试 | 系统正常运行 | 1. 输入不符合格式要求的数据\n2. 点击提交按钮 | 显示格式错误提示，数据不入库 | 高 |
| TC-GEN-006 | 性能测试 | 系统正常运行 | 1. 同时发起多个请求\n2. 记录响应时间 | 系统响应时间<2秒，无崩溃 | 中 |
| TC-GEN-007 | 并发测试 | 系统正常运行 | 1. 多用户同时操作同一数据\n2. 检查数据一致性 | 数据一致性保持，无脏数据 | 高 |
| TC-GEN-008 | 权限测试 | 系统正常运行 | 1. 使用无权限账号访问\n2. 尝试执行受限操作 | 提示无权限，操作被拒绝 | 高 |
| TC-GEN-009 | 数据持久化测试 | 操作已完成 | 1. 执行数据操作\n2. 重启系统\n3. 检查数据 | 数据正常保存，重启后不丢失 | 高 |
| TC-GEN-010 | 错误恢复测试 | 系统运行中 | 1. 模拟系统异常中断\n2. 重新启动系统\n3. 检查恢复情况 | 系统正常恢复，数据一致 | 中 |

如需下载测试用例，请点击消息下方的下载按钮！"""

    def _generate_code_example(self, message):
        return f"""[思考]用户需要代码示例或测试脚本。我需要提供一个通用的自动化测试脚本模板。

生成代码示例...[/思考]

## 自动化测试脚本示例（Python + pytest）

```python
import pytest
from playwright.sync_api import sync_playwright

class TestExample:
    \"\"\"示例测试类\"\"\"
    
    def setup_method(self):
        \"\"\"测试前置条件\"\"\"
        self.playwright = sync_playwright().start()
        self.browser = self.playwright.chromium.launch(headless=True)
        self.page = self.browser.new_page()
    
    def teardown_method(self):
        \"\"\"测试后置条件\"\"\"
        self.browser.close()
        self.playwright.stop()
    
    def test_login_success(self):
        \"\"\"测试登录成功流程\"\"\"
        # 打开登录页面
        self.page.goto("https://example.com/login")
        
        # 输入用户名和密码
        self.page.fill('input[name="username"]', "testuser")
        self.page.fill('input[name="password"]', "password123")
        
        # 点击登录按钮
        self.page.click('button[type="submit"]')
        
        # 验证登录成功
        assert self.page.url == "https://example.com/home"
        assert self.page.locator('.user-name').text_content() == "testuser"
    
    def test_login_with_invalid_password(self):
        \"\"\"测试使用错误密码登录\"\"\"
        self.page.goto("https://example.com/login")
        
        self.page.fill('input[name="username"]', "testuser")
        self.page.fill('input[name="password"]', "wrongpassword")
        self.page.click('button[type="submit"]')
        
        # 验证错误提示
        error_message = self.page.locator('.error-message').text_content()
        assert "用户名或密码错误" in error_message

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
```

**使用说明：**
1. 安装依赖：`pip install pytest playwright`
2. 初始化浏览器：`playwright install`
3. 运行测试：`pytest test_example.py -v`

如需下载代码，请点击消息下方的下载按钮！"""

    def _generate_document_example(self, message):
        return f"""[思考]用户需要文档或报告模板。我需要提供一个测试报告模板。

生成测试报告模板...[/思考]

## 测试报告模板

### 1. 测试概述

| 项目 | 内容 |
|-----|------|
| 测试项目 | [项目名称] |
| 测试周期 | [开始日期] ~ [结束日期] |
| 测试环境 | [环境描述] |
| 测试人员 | [测试人员] |

### 2. 测试执行统计

| 指标 | 数量 | 百分比 |
|-----|------|--------|
| 总用例数 | [总数] | 100% |
| 通过 | [通过数] | [通过率]% |
| 失败 | [失败数] | [失败率]% |
| 阻塞 | [阻塞数] | [阻塞率]% |

### 3. 缺陷统计

| 严重等级 | 数量 | 占比 |
|---------|------|------|
| 致命 | [数量] | [占比]% |
| 严重 | [数量] | [占比]% |
| 一般 | [数量] | [占比]% |
| 轻微 | [数量] | [占比]% |

### 4. 测试结论

- 测试覆盖度：[百分比]%
- 主要问题：[问题描述]
- 建议：[改进建议]

### 5. 附录

- 详细缺陷列表
- 测试用例执行明细
- 环境配置说明

如需下载报告模板，请点击消息下方的下载按钮！"""

    @action(detail=False, methods=['post'])
    def send_message_stream(self, request):
        session_id = request.data.get('session_id')
        message = (request.data.get('message') or '').strip()

        if not session_id or not message:
            return Response(
                {'error': 'session_id and message are required'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            session = AssistantSession.objects.get(session_id=session_id, user=request.user)
        except AssistantSession.DoesNotExist:
            return Response({'error': 'session not found'}, status=status.HTTP_404_NOT_FOUND)

        session.save(update_fields=['updated_at'])

        user_message = ChatMessage.objects.create(
            session=session,
            role='user',
            content=message,
            conversation_id=session.conversation_id,
        )

        temp_messages = list(ChatMessage.objects.filter(session=session).order_by('created_at'))
        temp_messages = [m for m in temp_messages if m.id != user_message.id]

        messages_for_ai = [{"role": "system", "content": self.SYSTEM_PROMPT}]
        for msg in temp_messages[-10:]:
            if msg.role == 'user':
                messages_for_ai.append({"role": "user", "content": msg.content})
            else:
                messages_for_ai.append({"role": "assistant", "content": msg.content})
        messages_for_ai.append({"role": "user", "content": message})
        
        logger.info(f"会话 {session_id} 构建了 {len(messages_for_ai)} 条上下文消息，包含 {len(temp_messages)} 条历史")

        q = queue.Queue()
        stop_event = threading.Event()
        full_content = {'value': ''}

        def async_worker():
            import asyncio
            from asgiref.sync import sync_to_async
            
            async def _run():
                try:
                    async for chunk in AIModelService.call_with_auto_model_stream(
                        messages_for_ai,
                        max_tokens=4096,
                        role='ai_tester'
                    ):
                        full_content['value'] += chunk
                        q.put(('content', chunk))
                        if stop_event.is_set():
                            break
                except Exception as e:
                    logger.error(f"流式AI调用失败: {e}")
                    mock_content = self._get_mock_response(message)
                    mock_content = self._append_appreciation(message, mock_content)
                    full_content['value'] = mock_content
                    chunk_size = 1
                    delay = 0.01
                    for i in range(0, len(mock_content), chunk_size):
                        chunk = mock_content[i:i+chunk_size]
                        q.put(('content', chunk))
                        if stop_event.is_set():
                            break
                        time.sleep(delay)
                    return
                
                appreciation = self.APPRECIATION_TEXT if self._should_add_appreciation(message) and '赞赏码' not in full_content['value'] else ''
                if appreciation:
                    for i in range(0, len(appreciation), 1):
                        chunk = appreciation[i:i+1]
                        full_content['value'] += chunk
                        q.put(('content', chunk))
                        if stop_event.is_set():
                            break
                        time.sleep(0.01)
                
                try:
                    create_message = sync_to_async(ChatMessage.objects.create)
                    assistant_message = await create_message(
                        session=session,
                        role='assistant',
                        content=full_content['value'],
                        conversation_id=session.conversation_id,
                    )
                    q.put(('complete', {
                        'user_message': ChatMessageSerializer(user_message).data,
                        'assistant_message': ChatMessageSerializer(assistant_message).data,
                        'model': '智能助手',
                        'usage': {'total_tokens': len(message) // 3 + len(full_content['value']) // 3}
                    }))
                except Exception as e:
                    logger.error(f"保存助手消息失败: {e}")
                
                q.put(('done', None))

            try:
                asyncio.run(_run())
            except Exception as e:
                logger.error(f"异步工作线程异常: {e}")
                q.put(('done', None))

        thread = threading.Thread(target=async_worker, daemon=True)
        thread.start()

        def sync_generator():
            while not stop_event.is_set():
                try:
                    item = q.get(timeout=1.0)
                    if item[0] == 'content':
                        yield f"data: {json.dumps({'type': 'content', 'content': item[1]})}\n\n"
                    elif item[0] == 'complete':
                        yield f"data: {json.dumps({'type': 'complete', **item[1]})}\n\n"
                    elif item[0] == 'done':
                        yield "data: [DONE]\n\n"
                        break
                except queue.Empty:
                    continue

        response = StreamingHttpResponse(sync_generator(), content_type='text/event-stream')
        response['Cache-Control'] = 'no-cache'
        response['Connection'] = 'keep-alive'
        return response

    @action(detail=False, methods=['post'])
    def download_document(self, request):
        session_id = request.data.get('session_id')
        message_id = request.data.get('message_id')
        format_type = request.data.get('format', 'word')

        if not session_id or not message_id:
            return Response(
                {'error': 'session_id and message_id are required'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            session = AssistantSession.objects.get(session_id=session_id, user=request.user)
            message = ChatMessage.objects.get(id=message_id, session=session)
        except (AssistantSession.DoesNotExist, ChatMessage.DoesNotExist):
            return Response({'error': 'session or message not found'}, status=status.HTTP_404_NOT_FOUND)

        content = message.content

        if format_type == 'word':
            return self._generate_word_document(content, message_id)
        elif format_type == 'pdf':
            return self._generate_pdf_document(content, message_id)
        elif format_type == 'csv':
            return self._generate_csv_document(content, message_id)
        else:
            return Response({'error': 'unsupported format'}, status=status.HTTP_400_BAD_REQUEST)

    def _generate_word_document(self, content, message_id):
        from docx import Document
        from docx.shared import Pt, Cm
        from io import BytesIO

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)

        lines = content.split('\n')
        in_table = False
        table_rows = []

        for line in lines:
            if line.startswith('|') and line.endswith('|'):
                in_table = True
                cells = [c.strip() for c in line.split('|')[1:-1]]
                table_rows.append(cells)
            else:
                if in_table and table_rows:
                    if len(table_rows) >= 2:
                        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                        for i, row in enumerate(table_rows):
                            for j, cell in enumerate(row):
                                table.cell(i, j).text = cell
                        for row in table.rows:
                            for cell in row.cells:
                                paragraphs = cell.paragraphs
                                for paragraph in paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)
                    in_table = False
                    table_rows = []
                if line.strip():
                    if line.startswith('# '):
                        heading = doc.add_heading(line[2:], level=1)
                        heading.style.font.name = '微软雅黑'
                    elif line.startswith('## '):
                        heading = doc.add_heading(line[3:], level=2)
                        heading.style.font.name = '微软雅黑'
                    elif line.startswith('### '):
                        heading = doc.add_heading(line[4:], level=3)
                        heading.style.font.name = '微软雅黑'
                    elif line.startswith('**') and line.endswith('**'):
                        para = doc.add_paragraph()
                        run = para.add_run(line[2:-2])
                        run.bold = True
                    else:
                        doc.add_paragraph(line)

        if in_table and table_rows and len(table_rows) >= 2:
            table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            for i, row in enumerate(table_rows):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = cell

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from django.http import HttpResponse
        response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="zx_ai_response_{message_id}.docx"'
        return response

    def _generate_pdf_document(self, content, message_id):
        from io import BytesIO
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        try:
            pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))
            chinese_font = 'SimHei'
        except:
            chinese_font = 'Helvetica'

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()

        style_normal = ParagraphStyle(
            'Normal',
            fontName=chinese_font,
            fontSize=10,
            leading=14,
        )
        style_heading1 = ParagraphStyle(
            'Heading1',
            fontName=chinese_font,
            fontSize=16,
            leading=20,
            bold=True,
        )
        style_heading2 = ParagraphStyle(
            'Heading2',
            fontName=chinese_font,
            fontSize=14,
            leading=18,
            bold=True,
        )

        elements = []
        lines = content.split('\n')
        in_table = False
        table_rows = []

        for line in lines:
            if line.startswith('|') and line.endswith('|'):
                in_table = True
                cells = [c.strip() for c in line.split('|')[1:-1]]
                table_rows.append(cells)
            else:
                if in_table and table_rows:
                    if len(table_rows) >= 2:
                        table = Table(table_rows)
                        table.setStyle(TableStyle([
                            ('FONTNAME', (0, 0), (-1, -1), chinese_font),
                            ('FONTSIZE', (0, 0), (-1, -1), 8),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ]))
                        elements.append(table)
                    in_table = False
                    table_rows = []
                if line.strip():
                    if line.startswith('# '):
                        elements.append(Paragraph(line[2:], style_heading1))
                    elif line.startswith('## '):
                        elements.append(Paragraph(line[3:], style_heading2))
                    elif line.startswith('**') and line.endswith('**'):
                        elements.append(Paragraph(f'<b>{line[2:-2]}</b>', style_normal))
                    else:
                        elements.append(Paragraph(line, style_normal))

        if in_table and table_rows and len(table_rows) >= 2:
            table = Table(table_rows)
            table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), chinese_font),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ]))
            elements.append(table)

        doc.build(elements)
        buffer.seek(0)

        from django.http import HttpResponse
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="zx_ai_response_{message_id}.pdf"'
        return response

    def _generate_csv_document(self, content, message_id):
        import csv
        from io import StringIO

        lines = content.split('\n')
        buffer = StringIO()
        writer = csv.writer(buffer)

        for line in lines:
            if line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                writer.writerow(cells)
            elif line.strip() and not line.startswith('#'):
                writer.writerow([line])

        buffer.seek(0)

        from django.http import HttpResponse
        response = HttpResponse(buffer.getvalue(), content_type='text/csv; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="zx_ai_response_{message_id}.csv"'
        response.write('\ufeff')
        return response


def assistant_view(request):
    """Legacy iframe assistant view."""
    return render(request, 'assistant/assistant.html')

from rest_framework import viewsets, permissions
from rest_framework.decorators import action
from rest_framework.response import Response
from django.db.models import Count, Q, Sum, F, Avg
from django.db.models.functions import TruncDate, Length
from django.utils import timezone
from datetime import timedelta, datetime
from .models import TestReport, ReportTemplate
from .serializers import TestReportSerializer, ReportTemplateSerializer
from apps.executions.models import TestPlan, TestRun, TestRunCase
from apps.testcases.models import TestCase
from apps.requirement_analysis.models import RequirementAnalysis, GeneratedTestCase, BusinessRequirement
import os
import logging

logger = logging.getLogger(__name__)


class TestReportViewSet(viewsets.ModelViewSet):
    """测试报告视图集"""
    queryset = TestReport.objects.all()
    serializer_class = TestReportSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        queryset = super().get_queryset()
        project_id = self.request.query_params.get('project')
        if project_id:
            queryset = queryset.filter(project_id=project_id)
        return queryset.order_by('-created_at')

    def perform_destroy(self, instance):
        """删除报告时清理关联的媒体文件"""
        # 清理报告关联的截图和视频目录
        try:
            from django.conf import settings
            import shutil

            # 尝试清理报告关联的生成文件目录
            report_dir = os.path.join(settings.BASE_DIR, 'media', 'reports', str(instance.id))
            if os.path.isdir(report_dir):
                shutil.rmtree(report_dir, ignore_errors=True)
                logger.info(f"已清理报告 #{instance.id} 的媒体文件目录: {report_dir}")
        except Exception as e:
            logger.warning(f"清理报告 #{instance.id} 文件时出错: {e}")
        instance.delete()

    @action(detail=False, methods=['post'], url_path='batch_delete')
    def batch_delete(self, request):
        """批量删除测试报告（含文件清理）"""
        ids = request.data.get('ids', [])
        if not ids:
            return Response({'error': '请提供要删除的报告 ID 列表'}, status=400)

        deleted = 0
        errors = []
        for report_id in ids:
            try:
                report = TestReport.objects.get(id=report_id)
                self.perform_destroy(report)
                deleted += 1
            except TestReport.DoesNotExist:
                errors.append(f'报告 #{report_id} 不存在')
            except Exception as e:
                errors.append(f'删除报告 #{report_id} 失败: {str(e)}')

        return Response({'deleted': deleted, 'errors': errors})
    
    @action(detail=False, methods=['get'])
    def dashboard(self, request):
        """获取概览数据"""
        try:
            project_id = request.query_params.get('project')

            # 基础查询集
            plans_qs = TestPlan.objects.filter(is_active=True)
            cases_qs = TestCase.objects.all()

            if project_id:
                plans_qs = plans_qs.filter(projects__id=project_id)
                cases_qs = cases_qs.filter(project_id=project_id)

            # 统计数据
            total_plans = plans_qs.count()
            total_cases = cases_qs.count()

            # 计算测试计划总进度
            # 遍历所有活跃计划，计算其下所有TestRun的进度平均值
            total_progress = 0
            plan_count_for_progress = 0

            for plan in plans_qs:
                runs = plan.test_runs.all()
                if runs.exists():
                    # 计算该计划下所有Run的平均进度
                    run_progresses = []
                    for run in runs:
                        try:
                            stats = run.progress_stats
                            if stats and isinstance(stats, dict) and 'progress' in stats:
                                run_progresses.append(stats['progress'])
                        except Exception:
                            pass
                    if run_progresses:
                        plan_progress = sum(run_progresses) / len(run_progresses)
                        total_progress += plan_progress
                        plan_count_for_progress += 1

            avg_plan_progress = round(total_progress / plan_count_for_progress, 1) if plan_count_for_progress > 0 else 0

            # 计算整体通过率
            recent_runs = TestRun.objects.filter(test_plan__in=plans_qs).order_by('-created_at')[:10]
            total_executed = 0
            total_passed = 0

            for run in recent_runs:
                try:
                    stats = run.progress_stats
                    if stats and isinstance(stats, dict):
                        total_executed += stats.get('tested', 0)
                        total_passed += stats.get('passed', 0)
                except Exception:
                    pass

            pass_rate = round((total_passed / total_executed * 100), 1) if total_executed > 0 else 0

            # 统计缺陷总数 (基于 TestRunCase 的 defects 字段)
            all_runs = TestRun.objects.filter(test_plan__in=plans_qs)
            defects_count = 0
            for run in all_runs:
                try:
                    run_cases_with_defects = run.run_cases.exclude(defects=[])
                    for rc in run_cases_with_defects:
                        if isinstance(rc.defects, list):
                            defects_count += len(rc.defects)
                except Exception:
                    pass

            return Response({
                'active_plans': total_plans,
                'plan_progress': avg_plan_progress,
                'total_cases': total_cases,
                'total_defects': defects_count,
                'pass_rate': pass_rate
            })
        except Exception as e:
            logger.error(f"获取 dashboard 数据失败: {str(e)}", exc_info=True)
            return Response({
                'active_plans': 0,
                'plan_progress': 0,
                'total_cases': 0,
                'total_defects': 0,
                'pass_rate': 0
            })

    @action(detail=False, methods=['get'])
    def status_distribution(self, request):
        """获取执行状态分布"""
        try:
            project_id = request.query_params.get('project')
            version_id = request.query_params.get('version')

            runs_qs = TestRun.objects.all()
            if project_id:
                runs_qs = runs_qs.filter(project_id=project_id)
            if version_id:
                runs_qs = runs_qs.filter(version_id=version_id)

            distribution = TestRunCase.objects.filter(test_run__in=runs_qs).values('status').annotate(
                count=Count('id')
            )

            result = {item['status']: item['count'] for item in distribution}
            for status, _ in TestRunCase.STATUS_CHOICES:
                if status not in result:
                    result[status] = 0

            return Response(result)
        except Exception as e:
            logger.error(f"获取 status_distribution 失败: {str(e)}", exc_info=True)
            result = {}
            for status, _ in TestRunCase.STATUS_CHOICES:
                result[status] = 0
            return Response(result)

    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """下载报告（支持 pdf、word、json 格式）"""
        from django.http import HttpResponse
        import json
        from rest_framework import status as http_status
        report = self.get_object()
        fmt = request.query_params.get('format', 'json')

        if fmt == 'json':
            data = {
                'report_name': report.name,
                'report_type': report.report_type,
                'created_at': report.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'summary': report.summary,
                'content': report.content
            }
            response = HttpResponse(json.dumps(data, ensure_ascii=False, indent=2), content_type='application/json')
            response['Content-Disposition'] = f'attachment; filename="{report.name}.json"'
            return response

        elif fmt == 'pdf':
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            buffer = HttpResponse(content_type='application/pdf')
            buffer['Content-Disposition'] = f'attachment; filename="{report.name}.pdf"'
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            story.append(Paragraph(f"<b>{report.name}</b>", styles['Title']))
            story.append(Paragraph(f"生成时间: {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Spacer(1, 12))
            summary = report.summary or {}
            story.append(Paragraph(f"总用例: {summary.get('total_cases', 0)} | 通过: {summary.get('passed', 0)} | 失败: {summary.get('failed', 0)} | 通过率: {summary.get('pass_rate', 0)}%", styles['Normal']))
            story.append(Spacer(1, 12))
            content = report.content or {}
            for tc in content.get('test_cases', []):
                story.append(Paragraph(f"<b>{tc.get('title', '用例')}</b> [{tc.get('status', '未知')}]", styles['Heading3']))
                story.append(Paragraph(f"步骤: {tc.get('steps', '无')}", styles['Normal']))
                story.append(Paragraph(f"预期: {tc.get('expected', '无')}", styles['Normal']))
                story.append(Spacer(1, 6))
            doc.build(story)
            return buffer

        elif fmt == 'word':
            from docx import Document
            doc = Document()
            doc.add_heading(report.name, 0)
            doc.add_paragraph(f"生成时间: {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            summary = report.summary or {}
            doc.add_paragraph(f"总用例: {summary.get('total_cases', 0)} | 通过: {summary.get('passed', 0)} | 失败: {summary.get('failed', 0)} | 通过率: {summary.get('pass_rate', 0)}%")
            content = report.content or {}
            for tc in content.get('test_cases', []):
                doc.add_heading(tc.get('title', '用例'), level=2)
                doc.add_paragraph(f"状态: {tc.get('status', '未知')}")
                doc.add_paragraph(f"测试步骤: {tc.get('steps', '无')}")
                doc.add_paragraph(f"预期结果: {tc.get('expected', '无')}")
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{report.name}.docx"'
            doc.save(response)
            return response

        else:
            return Response({'error': '不支持的格式，可选: json, pdf, word'}, status=http_status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['get'])
    def defect_distribution(self, request):
        """获取缺陷分布 (按优先级)"""
        try:
            project_id = request.query_params.get('project')
            qs = TestRunCase.objects.filter(status='failed')

            if project_id:
                qs = qs.filter(test_run__project_id=project_id)

            distribution = qs.values('priority').annotate(count=Count('id'))

            # 映射优先级显示
            priority_map = dict(TestRunCase.PRIORITY_CHOICES)
            result = []
            for item in distribution:
                result.append({
                    'name': priority_map.get(item['priority'], item['priority']),
                    'value': item['count']
                })

            return Response(result)
        except Exception as e:
            logger.error(f"获取 defect_distribution 失败: {str(e)}", exc_info=True)
            return Response([])

    @action(detail=False, methods=['get'])
    def failed_cases_top(self, request):
        """获取失败用例TOP榜"""
        try:
            project_id = request.query_params.get('project')

            qs = TestRunCase.objects.filter(status='failed')
            if project_id:
                qs = qs.filter(test_run__project_id=project_id)

            # 按 testcase 分组统计失败次数
            top_failed = qs.values(
                'testcase__id', 'testcase__title'
            ).annotate(
                fail_count=Count('id')
            ).order_by('-fail_count')[:10]

            return Response(top_failed)
        except Exception as e:
            logger.error(f"获取 failed_cases_top 失败: {str(e)}", exc_info=True)
            return Response([])

    @action(detail=False, methods=['get'])
    def execution_trend(self, request):
        """获取每日执行趋势"""
        try:
            project_id = request.query_params.get('project')
            days = int(request.query_params.get('days', 7))

            # 获取当前时区的今天开始时间
            current_tz = timezone.get_current_timezone()
            local_now = timezone.localtime(timezone.now())
            today = local_now.date()

            # 计算起始日期
            start_date = today - timedelta(days=days - 1)

            # 构造起始时间的 datetime 对象 (00:00:00)
            start_datetime = datetime.combine(start_date, datetime.min.time())
            start_datetime = timezone.make_aware(start_datetime, current_tz)

            qs = TestRunCase.objects.filter(
                executed_at__gte=start_datetime,
                status__in=['passed', 'failed', 'blocked', 'retest']
            )

            if project_id:
                qs = qs.filter(test_run__project_id=project_id)

            # 由于数据库聚合(TruncDate)在某些环境下返回None，改为Python内存聚合
            # 获取所有符合条件的记录的执行时间
            executions = qs.values_list('executed_at', flat=True)

            # 初始化日期映射
            date_map = {}

            for executed_at in executions:
                if executed_at:
                    # 转换为本地时间
                    local_time = executed_at.astimezone(current_tz)
                    date_str = local_time.date().strftime('%Y-%m-%d')
                    date_map[date_str] = date_map.get(date_str, 0) + 1

            # 补全日期
            result = []
            for i in range(days):
                date = start_date + timedelta(days=i)
                date_str = date.strftime('%Y-%m-%d')
                result.append({
                    'date': date_str,
                    'count': date_map.get(date_str, 0)
                })

            return Response(result)
        except Exception as e:
            logger.error(f"获取 execution_trend 失败: {str(e)}", exc_info=True)
            return Response([])

    @action(detail=False, methods=['get'])
    def ai_efficiency(self, request):
        """获取AI效能分析"""
        try:
            project_id = request.query_params.get('project')

            cases_qs = TestCase.objects.all()
            generated_qs = GeneratedTestCase.objects.all()
            requirements_qs = BusinessRequirement.objects.all()

            if project_id:
                cases_qs = cases_qs.filter(project_id=project_id)
                generated_qs = generated_qs.filter(requirement__analysis__document__project_id=project_id)
                requirements_qs = requirements_qs.filter(analysis__document__project_id=project_id)

            # 1. AI生成 vs 人工创建
            ai_count = generated_qs.count()
            adopted_ai_count = generated_qs.filter(status='adopted').count()
            total_cases = cases_qs.count()
            manual_count = max(0, total_cases - adopted_ai_count)

            # 2. 生成采纳率
            adoption_rate = round((adopted_ai_count / ai_count * 100), 1) if ai_count > 0 else 0

            # 3. 需求覆盖率
            total_reqs = requirements_qs.count()
            covered_reqs = generated_qs.filter(status='adopted').values('requirement').distinct().count()
            coverage_rate = round((covered_reqs / total_reqs * 100), 1) if total_reqs > 0 else 0

            # 4. 节省时间估算
            saved_hours = round(ai_count * 15 / 60, 1)

            return Response({
                'ai_vs_manual': {
                    'ai': ai_count,
                    'manual': manual_count
                },
                'adoption_rate': adoption_rate,
                'requirement_coverage': coverage_rate,
                'saved_hours': saved_hours
            })
        except Exception as e:
            logger.error(f"获取 ai_efficiency 失败: {str(e)}", exc_info=True)
            return Response({
                'ai_vs_manual': {'ai': 0, 'manual': 0},
                'adoption_rate': 0,
                'requirement_coverage': 0,
                'saved_hours': 0
            })

    @action(detail=False, methods=['get'])
    def team_workload(self, request):
        """获取团队工作量"""
        try:
            project_id = request.query_params.get('project')

            qs = TestRunCase.objects.filter(
                status__in=['passed', 'failed', 'blocked', 'retest'],
                executed_by__isnull=False
            )

            if project_id:
                qs = qs.filter(test_run__project_id=project_id)

            # 统计执行数量
            execution_stats = qs.values(
                'executed_by__username'
            ).annotate(
                count=Count('id')
            ).order_by('-count')[:10]

            # 统计发现缺陷数量
            defect_stats = {}
            defect_qs = qs.filter(status__in=['failed', 'blocked'])
            defect_data = defect_qs.values('executed_by__username').annotate(count=Count('id'))
            for item in defect_data:
                defect_stats[item['executed_by__username']] = item['count']

            result = []
            for item in execution_stats:
                username = item['executed_by__username']
                result.append({
                    'username': username,
                    'execution_count': item['count'],
                    'defect_count': defect_stats.get(username, 0)
                })

            return Response(result)
        except Exception as e:
            logger.error(f"获取 team_workload 失败: {str(e)}", exc_info=True)
            return Response([])
